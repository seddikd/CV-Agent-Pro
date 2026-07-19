from pathlib import Path
import logging

from pypdf import PdfReader
from docx import Document


log = logging.getLogger(__name__)


def extract_text(file_path: Path, max_chars: int = 30000) -> str:
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        text = _extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        text = _extract_docx(file_path)
    else:
        log.warning("Unsupported extension: %s", ext)
        return ""

    text = text.strip()
    if len(text) > max_chars:
        text = text[:max_chars] + "\n[...TRONQUÉ...]"
    return text


def _extract_pdf(path: Path) -> str:
    try:
        reader = PdfReader(str(path))
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n".join(parts)
    except Exception as e:
        log.warning("PDF extract failed for %s: %s", path.name, e)
        return ""


def _table_texts(table) -> list[str]:
    """Texte de toutes les cellules d'un tableau (tableaux imbriqués compris)."""
    out: list[str] = []
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if para.text.strip():
                    out.append(para.text)
            for nested in cell.tables:  # tableaux imbriqués (mise en page fréquente)
                out.extend(_table_texts(nested))
    return out


def _extract_docx(path: Path) -> str:
    """Texte d'un .docx : paragraphes + TABLEAUX + en-têtes/pieds de page.

    Beaucoup de CV Word sont mis en page dans des tableaux (ou placent le
    nom/contact dans l'en-tête). Ne lire que `doc.paragraphs` — comme avant —
    perdait tout ce contenu et faisait passer le CV pour un non-CV.
    """
    try:
        doc = Document(str(path))
    except Exception as e:
        log.warning("DOCX extract failed for %s: %s", path.name, e)
        return ""

    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        parts.extend(_table_texts(table))
    # En-têtes / pieds de page (nom, coordonnées y sont souvent placés).
    for section in doc.sections:
        for zone in (section.header, section.footer):
            if zone is not None:
                parts.extend(p.text for p in zone.paragraphs)

    # Déduplication en préservant l'ordre : les cellules fusionnées font répéter
    # le même objet cellule d'une ligne à l'autre (python-docx), d'où des doublons.
    vues: set[str] = set()
    lignes: list[str] = []
    for t in parts:
        t = t.strip()
        if t and t not in vues:
            vues.add(t)
            lignes.append(t)
    return "\n".join(lignes)
